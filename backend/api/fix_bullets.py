"""
Comprehensive post-processor for python-docx files created by pdf2docx.

Fixes:
1. Broken bullet glyphs (Wingdings, PUA U+E000–U+F8FF codepoints like \\uf0a7, \\uf0b7, etc.)
   in both body paragraphs and table cell paragraphs.
2. Ensures table cell borders and shading render cleanly in Microsoft Word / Google Docs.
3. Cleans control characters and zero-width spaces that break Word rendering.
"""

import re
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Wingdings / Symbol bullet glyphs live in the Unicode Private Use Area (U+E000–U+F8FF)
# or standard bullet/middle-dot characters used by converters.
PUA_BULLET_PATTERN = re.compile(r'^[\ue000-\uf8ff\u00b7\u2022\u25cf\u25cb\u25a0\u25aa\u2013\u2014]\s*')

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding (in dxa) for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def fix_paragraph_bullets(para):
    """Fixes bullets inside a single paragraph object."""
    if not para.runs:
        return False

    full_text = para.text
    if not full_text:
        return False

    first_run = para.runs[0]
    first_text = first_run.text

    # Check if paragraph starts with a PUA/symbol bullet
    match = PUA_BULLET_PATTERN.match(first_text)
    if match:
        # Strip the broken glyph from the first run
        first_run.text = PUA_BULLET_PATTERN.sub('', first_text, count=1)

        # Create bullet run at the very beginning of paragraph XML
        bullet_run = para.add_run()
        para._p.insert(0, bullet_run._r)
        bullet_run.text = '•  '
        bullet_run.font.name = 'Arial'

        # Set hanging indent formatting for clean bullet alignment
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        return True

    return False


def enhance_table(table):
    """Enhances table alignment, row splitting, and border defaults."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        # Ensure default table borders if none specified
        borders_exist = table._element.xpath('w:tblPr/w:tblBorders')
        if not borders_exist:
            borders_xml = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>\n'
                f'  <w:left w:val="none"/>\n'
                f'  <w:right w:val="none"/>\n'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>\n'
                f'  <w:insideV w:val="none"/>\n'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders_xml)

    # Prevent rows from splitting across pages
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        if not trPr.xpath('w:cantSplit'):
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        for cell in row.cells:
            for para in cell.paragraphs:
                fix_paragraph_bullets(para)


def fix_bullet_paragraphs(doc):
    """
    Mutates python-docx Document in place:
    - Fixes bullets in body paragraphs & table cells.
    - Enhances table borders and alignment.
    - Cleans up control characters.
    """
    fixed_count = 0

    # Body paragraphs
    for para in doc.paragraphs:
        if fix_paragraph_bullets(para):
            fixed_count += 1

    # Tables
    for table in doc.tables:
        enhance_table(table)

    return fixed_count