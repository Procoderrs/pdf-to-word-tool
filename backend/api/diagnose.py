import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import io
import sys

FLAG_ITALIC = 2
FLAG_BOLD = 16

POINTS_PER_INCH = 72


def convert_pdf_to_docx(pdf_path, docx_path):
    pdf = fitz.open(pdf_path)
    document = Document()

    for page_num, page in enumerate(pdf):
        text_dict = page.get_text("dict")

        for block in text_dict["blocks"]:
            if block["type"] == 0:
                # --- Text block: har PDF-paragraph = ek docx-paragraph ---
                paragraph = document.add_paragraph()

                for line_index, line in enumerate(block["lines"]):
                    for span in line["spans"]:
                        text = span["text"]
                        if not text.strip():
                            continue

                        run = paragraph.add_run(text)
                        run.font.size = Pt(round(span["size"]))

                        flags = span["flags"]
                        run.bold = bool(flags & FLAG_BOLD)
                        run.italic = bool(flags & FLAG_ITALIC)

                        # Color chahe kuch bhi ho, hamesha editable TEXT ki
                        # tarah hi rahega — kabhi image nahi banega.
                        color_int = span.get("color", 0)
                        r = (color_int >> 16) & 255
                        g = (color_int >> 8) & 255
                        b = color_int & 255
                        run.font.color.rgb = RGBColor(r, g, b)

                    if line_index < len(block["lines"]) - 1:
                        paragraph.add_run(" ")

            elif block["type"] == 1:
                # --- Image block: asal image ki tarah hi docx mein daalo ---
                image_bytes = block.get("image")
                if not image_bytes:
                    continue

                bbox = block["bbox"]  # (x0, y0, x1, y1) in points
                width_pt = bbox[2] - bbox[0]
                width_in = width_pt / POINTS_PER_INCH

                try:
                    image_stream = io.BytesIO(image_bytes)
                    document.add_picture(image_stream, width=Inches(width_in))
                except Exception as e:
                    # Corrupt/unsupported image data — is image ko skip karo,
                    # baaki page ka conversion na ruke.
                    print(f"Image skip (page {page_num + 1}): {e}")

        if page_num < len(pdf) - 1:
            document.add_page_break()

    document.save(docx_path)
    pdf.close()


if __name__ == "__main__":
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    convert_pdf_to_docx(pdf_path, docx_path)
    print("DONE")