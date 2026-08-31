"""
Table Handler — turns PDF table-detection output into DOCX tables.
"""
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, RGBColor

TABLE_OVERLAP_THRESHOLD = 0.5


def set_repeat_header_row(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def set_cell_background(cell, rgb):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*rgb))
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_border(cell, edge, color="000000", sz=8):
    """edge is one of: top, bottom, left, right"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    el = OxmlElement(f'w:{edge}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), str(sz))
    el.set(qn('w:color'), color)
    tcBorders.append(el)


def is_dark_color(rgb):
    r, g, b = rgb
    luminance = 0.299 * r + 0.587 * g + 0.114 * b
    return luminance < 128


def get_row_fill_color(page, row_bbox):
    x0, y0, x1, y1 = row_bbox
    row_area = (x1 - x0) * (y1 - y0)
    if row_area <= 0:
        return None

    # FIX: header backgrounds are sometimes drawn as several small
    # rectangles (one per cell) rather than one big rect. Sum up total
    # overlap area PER COLOR instead of looking for a single rect that
    # covers most of the row — then pick whichever color covers the
    # most cumulative area.
    color_totals = {}
    for d in page.get_drawings():
        fill = d.get("fill")
        if not fill:
            continue
        for item in d["items"]:
            if item[0] != "re":
                continue
            rx0, ry0, rx1, ry1 = item[1]
            ix0, iy0 = max(x0, rx0), max(y0, ry0)
            ix1, iy1 = min(x1, rx1), min(y1, ry1)
            if ix1 <= ix0 or iy1 <= iy0:
                continue
            overlap = (ix1 - ix0) * (iy1 - iy0)
            key = tuple(round(c, 3) for c in fill)
            color_totals[key] = color_totals.get(key, 0) + overlap

    if not color_totals:
        return None

    best_color, best_total = max(color_totals.items(), key=lambda kv: kv[1])
    if best_total < row_area * 0.3:
        return None

    r, g, b = [int(c * 255) for c in best_color]
    return (r, g, b)

def add_table_to_doc(document, table_data, usable_width_in, page=None, table_obj=None):
    if not table_data or not table_data[0]:
        return

    num_rows = len(table_data)
    num_cols = len(table_data[0])

    docx_table = document.add_table(rows=num_rows, cols=num_cols)
    # FIX: no default 'Table Grid' style — that drew a full boxy grid
    # around every cell, including near-empty totals rows that should
    # look borderless in a typical invoice. We add only the borders we
    # actually want (header underline + outer table rules) below.
    docx_table.style = None
    docx_table.autofit = False
    col_width = Inches(usable_width_in / num_cols)
    for col in docx_table.columns:
        col.width = col_width

    # Detect the header row's real fill color from the PDF, if we have
    # enough info (page + the original fitz table object with row bboxes).
    header_fill = None
    print(f"[table_handler] page is None: {page is None}, table_obj is None: {table_obj is None}")
    if table_obj is not None:
        print(f"[table_handler] table_obj has 'rows': {hasattr(table_obj, 'rows')}")
        if hasattr(table_obj, "rows"):
            print(f"[table_handler] num rows: {len(table_obj.rows) if table_obj.rows else 0}")

    if page is not None and table_obj is not None and getattr(table_obj, "rows", None):
        try:
            first_row = table_obj.rows[0]
            print(f"[table_handler] first_row.cells: {first_row.cells}")
            cell_bboxes = [c for c in first_row.cells if c]
            if cell_bboxes:
                x0 = min(c[0] for c in cell_bboxes)
                y0 = min(c[1] for c in cell_bboxes)
                x1 = max(c[2] for c in cell_bboxes)
                y1 = max(c[3] for c in cell_bboxes)
                print(f"[table_handler] header row bbox: {(x0, y0, x1, y1)}")
                header_fill = get_row_fill_color(page, (x0, y0, x1, y1))
                print(f"[table_handler] header_fill result: {header_fill}")

                # extra debug: dump all drawings on this page so we can see
                # what fills actually exist, in case get_row_fill_color's
                # overlap threshold is too strict
                for i, d in enumerate(page.get_drawings()):
                    if d.get("fill"):
                        print(f"[table_handler]   drawing {i} fill={d['fill']}")
        except Exception as e:
            print(f"[table_handler] header fill detection skipped: {e}")
    for row_index, row in enumerate(table_data):
        for col_index, cell_text in enumerate(row):
            if col_index >= num_cols:
                continue
            cell = docx_table.cell(row_index, col_index)
            cell.text = str(cell_text) if cell_text is not None else ''
            cell.width = col_width

        if row_index == 0:
            set_repeat_header_row(docx_table.rows[0])

            if header_fill:
                use_white_text = is_dark_color(header_fill)
                for col_index in range(num_cols):
                    cell = docx_table.cell(0, col_index)
                    set_cell_background(cell, header_fill)
                    if use_white_text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)

            # underline under the header row
            for cell in docx_table.rows[0].cells:
                set_cell_border(cell, 'bottom')

    # outer table rules: a line above the header and a line below the
    # last row — no internal grid lines
    for cell in docx_table.rows[0].cells:
        set_cell_border(cell, 'top')
    for cell in docx_table.rows[-1].cells:
        set_cell_border(cell, 'bottom')