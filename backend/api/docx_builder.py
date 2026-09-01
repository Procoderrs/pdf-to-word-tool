"""
DOCX Engine — Stage 4/5 of the pipeline: takes parsed PDF pages and
builds the final .docx (paragraphs, runs, tables, images), applying
layout ordering from layout_parser.
"""
import io
import fitz
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .ocr_handler import get_page_text_dict
from .layout_parser import (
    clean_font_name, get_heading_style, enable_hyphenation,
    rects_overlap_ratio, spans_need_space, set_run_font,
    find_column_gap, order_items_for_columns, get_sidebar_fill_color,
    get_page_background_color, detect_list_type, get_horizontal_lines, split_line_by_large_gap,
    merge_same_line_blocks, merge_overlapping_lines,
    FULL_WIDTH_RATIO_THRESHOLD,
)
from .table_handler import add_table_to_doc

FLAG_ITALIC = 2
FLAG_BOLD = 16
POINTS_PER_INCH = 72
TABLE_OVERLAP_THRESHOLD = 0.5
CENTER_TOLERANCE_PT = 15  # how close a block's center must be to the content center to count as "centered"
NARROW_BLOCK_RATIO = 0.85  # a block must be narrower than this fraction of content width to be eligible for centering
ROW_PAIR_MIN_Y_OVERLAP_RATIO = 0.5  # how much two blocks' y-ranges must overlap to count as "same row"
COLUMN_GAP_STRADDLE_TOLERANCE_PT = 4  # slack when checking if a pair sits on either side of the real column gap
FULL_WIDTH_BULLET_INDENT_PT = 36  # hanging indent for bullets rendered at full page width (single column)
NARROW_CONTAINER_BULLET_INDENT_PT = 18  # FIX: smaller indent for bullets rendered inside a narrow 2-column table cell —
                                          # 36pt was eating too much of a narrow cell's width, squeezing/wrapping bullet
                                          # text badly (this is what broke bullets in 2-column layouts like the CV
                                          # sidebar, recipe ingredients list, etc. after the earlier 36pt change)


def detect_page_margins(pdf, page_width_pt):
    """Scan all pages' text blocks to find the real left/right content
    boundary, instead of assuming a fixed 1-inch margin. Returns the
    margins in inches plus the raw (min_x0, max_x1) content bounds in
    points, which callers use for width-aware centering decisions."""
    min_x0, max_x1 = page_width_pt, 0
    for page in pdf:
        text_dict = page.get_text("dict")
        for block in text_dict["blocks"]:
            if block["type"] == 0:
                min_x0 = min(min_x0, block["bbox"][0])
                max_x1 = max(max_x1, block["bbox"][2])
    if min_x0 >= max_x1:
        # no text found anywhere — fall back to a sane default
        return 1.0, 1.0, (0, page_width_pt)
    left_margin_in = max(min_x0 / POINTS_PER_INCH, 0.4)
    right_margin_in = max((page_width_pt - max_x1) / POINTS_PER_INCH, 0.4)
    return left_margin_in, right_margin_in, (min_x0, max_x1)


def set_document_background(document, rgb):
    """Sets Word's page background color (w:background) and ensures it
    actually displays on-screen (w:displayBackgroundShape in settings —
    without this, Word only shows the background when printing/exporting,
    not while viewing). NOTE: this is a document-wide setting in the
    .docx format — Word has no concept of a different background per
    page, so if a PDF's pages have different backgrounds, only one
    (whichever is passed in) can be applied to the whole document."""
    hex_color = '{:02X}{:02X}{:02X}'.format(*rgb)

    background = OxmlElement('w:background')
    background.set(qn('w:color'), hex_color)
    document.element.insert(0, background)

    settings_element = document.settings.element
    display_bg = OxmlElement('w:displayBackgroundShape')
    settings_element.insert(0, display_bg)


def write_span_to_paragraph(para, span, prev_span):
    """MOVED: was a nested closure inside add_block_to_container. Now
    module-level so add_row_pair_to_container (below) can reuse it too."""
    text = span["text"]
    if not text.strip():
        return prev_span
    if spans_need_space(prev_span, span):
        para.add_run(" ")
    run = para.add_run(text)
    run.font.size = Pt(round(span["size"]))
    set_run_font(run, clean_font_name(span.get("font", "")))
    flags = span["flags"]
    run.bold = bool(flags & FLAG_BOLD)
    run.italic = bool(flags & FLAG_ITALIC)
    color_int = span.get("color", 0)
    r = (color_int >> 16) & 255
    g = (color_int >> 8) & 255
    b = color_int & 255
    run.font.color.rgb = RGBColor(r, g, b)
    return span


def _pair_straddles_column_gap(left_bbox, right_bbox, column_gap):
    """Same guard as layout_parser._pair_straddles_column_gap — kept
    local here too since find_row_pairs lives in this module."""
    if not column_gap:
        return False
    gap_start, gap_end = column_gap
    return (
        left_bbox[2] <= gap_start + COLUMN_GAP_STRADDLE_TOLERANCE_PT
        and right_bbox[0] >= gap_end - COLUMN_GAP_STRADDLE_TOLERANCE_PT
    )


def find_row_pairs(items, page_width_pt, column_gap=None):
    """Detects two separate text BLOCKS (not spans within one line)
    that sit on the same horizontal row — e.g. 'Johary — Jewellery...'
    and 'Live Site | Admin Panel | GitHub', or 'WODWES LLC...' and
    'September 2024 · Present | Faisalabad'. PyMuPDF sometimes extracts
    these as two independent blocks instead of one block with multiple
    spans on one line, so split_line_by_large_gap (which only compares
    spans inside a single line) never sees them together and the
    right-alignment tab-stop code never triggers.

    column_gap (gap_start, gap_end), if given, is the REAL two-column
    split already detected on this page BEFORE this function runs (see
    convert_pdf_to_docx). Any candidate pair that straddles that gap is
    skipped so genuine two-column layouts are never mistaken for a
    same-line split.
    """
    text_blocks = [
        (idx, item) for idx, item in enumerate(items)
        if item['kind'] == 'block' and item['data']['type'] == 0
    ]
    used = set()
    merged = []

    for a in range(len(text_blocks)):
        i, item_a = text_blocks[a]
        if i in used:
            continue
        a_bbox = item_a['bbox']
        for b in range(a + 1, len(text_blocks)):
            j, item_b = text_blocks[b]
            if j in used:
                continue
            b_bbox = item_b['bbox']

            y_overlap = min(a_bbox[3], b_bbox[3]) - max(a_bbox[1], b_bbox[1])
            min_height = min(a_bbox[3] - a_bbox[1], b_bbox[3] - b_bbox[1])
            if min_height <= 0 or y_overlap / min_height < ROW_PAIR_MIN_Y_OVERLAP_RATIO:
                continue

            if a_bbox[0] < b_bbox[0]:
                left_item, right_item = item_a, item_b
            else:
                left_item, right_item = item_b, item_a

            gap = right_item['bbox'][0] - left_item['bbox'][2]
            if gap <= 0:
                continue

            if _pair_straddles_column_gap(left_item['bbox'], right_item['bbox'], column_gap):
                continue

            used.add(i)
            used.add(j)
            merged.append({
                'kind': 'row_pair',
                'bbox': (
                    min(left_item['bbox'][0], right_item['bbox'][0]),
                    min(left_item['bbox'][1], right_item['bbox'][1]),
                    max(left_item['bbox'][2], right_item['bbox'][2]),
                    max(left_item['bbox'][3], right_item['bbox'][3]),
                ),
                'left_data': left_item['data'],
                'right_data': right_item['data'],
            })
            break

    if not merged:
        return items

    result = [item for idx, item in enumerate(items) if idx not in used]
    result.extend(merged)
    result.sort(key=lambda it: it['bbox'][1])
    return result


def add_row_pair_to_container(container, item, content_bounds):
    """Renders a merged row_pair (see find_row_pairs) as a single
    paragraph — left block's text, a tab, then right block's text —
    with a right-aligned tab-stop at the content's right edge."""
    paragraph = container.add_paragraph()

    if content_bounds:
        min_x0, max_x1 = content_bounds
        right_tab_pos_in = (max_x1 - min_x0) / POINTS_PER_INCH
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(right_tab_pos_in), WD_TAB_ALIGNMENT.RIGHT
        )

    prev_span = None
    for line in item['left_data']['lines']:
        for span in line['spans']:
            prev_span = write_span_to_paragraph(paragraph, span, prev_span)

    paragraph.add_run("\t")
    prev_span = None
    for line in item['right_data']['lines']:
        for span in line['spans']:
            prev_span = write_span_to_paragraph(paragraph, span, prev_span)

    return paragraph


def add_block_to_container(container, block, page_width_pt=None, content_bounds=None,
                            horizontal_lines=None, narrow_container=False):
    """container = document OR a table cell — both support add_paragraph().
    page_width_pt / content_bounds are optional; when given, enables
    content-width-aware center-alignment detection for the block.

    narrow_container: FIX — set True when this block is being rendered
    inside a 2-column table cell (see render_two_column_table below),
    so bullet hanging-indent uses a smaller value than at full page
    width. A single fixed 36pt indent worked fine for full-width single-
    column pages but ate too much of a narrow cell's width, squeezing
    or badly wrapping bullet text in every 2-column layout (CV sidebar,
    recipe ingredients column, etc.)."""
    if block["type"] == 0:
        lines = block["lines"]
        lines = merge_overlapping_lines(lines)
        line_gaps = []
        for i in range(1, len(lines)):
            gap = lines[i]["bbox"][1] - lines[i - 1]["bbox"][3]
            line_gaps.append(gap)
        line_gaps.sort()

        if len(line_gaps) >= 4:
            typical_gap = line_gaps[len(line_gaps) // 2]
        elif len(line_gaps) >= 2:
            typical_gap = min(line_gaps)
        else:
            typical_gap = 0
        split_threshold = typical_gap + 0.19

        paragraph = container.add_paragraph()
        max_size_in_paragraph = 0
        prev_span = None
        prev_line_bbox = None

        bullet_indent_pt = (
            NARROW_CONTAINER_BULLET_INDENT_PT if narrow_container else FULL_WIDTH_BULLET_INDENT_PT
        )

        def apply_heading_style_to(para, size):
            if not para.runs:
                return
            first_run = para.runs[0]
            list_type, cleaned_text = detect_list_type(first_run.text)
            if list_type:
                para.style = list_type
                first_run.text = cleaned_text
                para.paragraph_format.left_indent = Pt(bullet_indent_pt)
                para.paragraph_format.first_line_indent = Pt(-bullet_indent_pt)
                para.paragraph_format.space_after = Pt(2)

        def line_max_size(line):
            sizes = [s["size"] for s in line["spans"] if s["text"].strip()]
            return max(sizes) if sizes else 0

        for line_index, line in enumerate(lines):
            line_text_preview = "".join(s["text"] for s in line["spans"]).strip()
            list_type_preview, _ = detect_list_type(line_text_preview)

            left_spans, right_spans = split_line_by_large_gap(line)

            if prev_line_bbox is not None:
                gap = line["bbox"][1] - prev_line_bbox[3]
                gap_split = gap >= split_threshold
                bullet_split = list_type_preview is not None

                prev_size = line_max_size(lines[line_index - 1])
                curr_size = line_max_size(line)
                size_split = (
                    prev_size and curr_size
                    and abs(curr_size - prev_size) / prev_size > 0.15
                )

                will_split = gap_split or bullet_split or size_split

                if will_split:
                    apply_heading_style_to(paragraph, max_size_in_paragraph)
                    paragraph.paragraph_format.space_after = Pt(round(max(0, min(gap, 12))))
                    paragraph = container.add_paragraph()
                    max_size_in_paragraph = 0
                    prev_span = None
                else:
                    paragraph.add_run(" ")

            if right_spans and content_bounds:
                min_x0, max_x1 = content_bounds
                right_tab_pos_in = (max_x1 - min_x0) / POINTS_PER_INCH
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(right_tab_pos_in),
                    WD_TAB_ALIGNMENT.RIGHT
                )
                for span in left_spans:
                    prev_span = write_span_to_paragraph(paragraph, span, prev_span)
                    max_size_in_paragraph = max(max_size_in_paragraph, span["size"])
                paragraph.add_run("\t")
                prev_span = None
                for span in right_spans:
                    prev_span = write_span_to_paragraph(paragraph, span, prev_span)
                    max_size_in_paragraph = max(max_size_in_paragraph, span["size"])
            else:
                for span in left_spans:
                    prev_span = write_span_to_paragraph(paragraph, span, prev_span)
                    max_size_in_paragraph = max(max_size_in_paragraph, span["size"])

            prev_line_bbox = line["bbox"]

        apply_heading_style_to(paragraph, max_size_in_paragraph)

        if horizontal_lines and content_bounds:
            min_x0, max_x1 = content_bounds
            content_width = max_x1 - min_x0
            if find_line_below(block["bbox"], horizontal_lines, content_width * 0.7):
                set_paragraph_bottom_border(paragraph)

        if content_bounds:
            min_x0, max_x1 = content_bounds
            content_width = max_x1 - min_x0
            content_center = (min_x0 + max_x1) / 2

            block_bbox = block["bbox"]
            block_width = block_bbox[2] - block_bbox[0]
            block_center = (block_bbox[0] + block_bbox[2]) / 2

            block_text = " ".join(
                span["text"] for line in lines for span in line["spans"]
            ).strip()
            is_page_number = block_text.isdigit() and len(block_text) <= 4

            is_narrow_and_centered = (
                block_width < content_width * NARROW_BLOCK_RATIO
                and abs(block_center - content_center) < CENTER_TOLERANCE_PT
            )

            if is_page_number or is_narrow_and_centered:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif page_width_pt:
            block_bbox = block["bbox"]
            block_center = (block_bbox[0] + block_bbox[2]) / 2
            page_center = page_width_pt / 2
            if abs(block_center - page_center) < CENTER_TOLERANCE_PT:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif block["type"] == 1:
        image_bytes = block.get("image")
        if not image_bytes:
            return
        bbox = block["bbox"]
        width_in = (bbox[2] - bbox[0]) / POINTS_PER_INCH

        try:
            add_picture_to_container(container, io.BytesIO(image_bytes), width_in)
            return
        except Exception as e:
            print(f"Image skip (direct): {type(e).__name__}: {e!r}")

        try:
            from PIL import Image
            img = Image.open(io.BytesIO(image_bytes))
            if img.mode not in ("RGB", "RGBA"):
                img = img.convert("RGB")
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            buf.seek(0)
            add_picture_to_container(container, buf, width_in)
        except Exception as e2:
            print(f"Image skip (PIL fallback also failed): {type(e2).__name__}: {e2!r}")


def set_cell_background(cell, rgb):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*rgb))
    cell._tc.get_or_add_tcPr().append(shd)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Values are in twentieths-of-a-point (dxa). 100 ≈ 0.07 inch."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)


def render_two_column_table(document, left_items, right_items, gap_start, gap_end,
                             page_width_pt, usable_width_in, sidebar_color):
    left_width_in = (gap_start / page_width_pt) * usable_width_in
    gutter_width_in = ((gap_end - gap_start) / page_width_pt) * usable_width_in
    right_width_in = usable_width_in - left_width_in - gutter_width_in

    has_gutter = gutter_width_in > 0.02
    num_cols = 3 if has_gutter else 2

    table = document.add_table(rows=1, cols=num_cols)
    table.autofit = False
    remove_table_borders(table)

    if has_gutter:
        table.columns[0].width = Inches(left_width_in)
        table.columns[1].width = Inches(gutter_width_in)
        table.columns[2].width = Inches(right_width_in)
        left_cell, spacer_cell, right_cell = table.rows[0].cells
        left_cell.width = Inches(left_width_in)
        spacer_cell.width = Inches(gutter_width_in)
        right_cell.width = Inches(right_width_in)
        set_cell_margins(spacer_cell, top=0, start=0, bottom=0, end=0)
        spacer_cell.paragraphs[0].text = ""
    else:
        table.columns[0].width = Inches(left_width_in)
        table.columns[1].width = Inches(right_width_in)
        left_cell, right_cell = table.rows[0].cells
        left_cell.width = Inches(left_width_in)
        right_cell.width = Inches(right_width_in)

    set_cell_margins(left_cell, top=100, start=100, bottom=100, end=50)
    set_cell_margins(right_cell, top=100, start=50, bottom=100, end=100)

    if sidebar_color:
        set_cell_background(left_cell, sidebar_color)

    left_cell.paragraphs[0].text = ""
    right_cell.paragraphs[0].text = ""

    # FIX: narrow_container=True — these cells are only a fraction of the
    # page width, so bullets here need the smaller indent (see
    # NARROW_CONTAINER_BULLET_INDENT_PT above).
    for item in left_items:
        if item['kind'] == 'row_pair':
            add_row_pair_to_container(left_cell, item, None)
        elif item['kind'] == 'block':
            add_block_to_container(left_cell, item['data'], narrow_container=True)
    for item in right_items:
        if item['kind'] == 'row_pair':
            add_row_pair_to_container(right_cell, item, None)
        elif item['kind'] == 'block':
            add_block_to_container(right_cell, item['data'], narrow_container=True)


def convert_pdf_to_docx(pdf_path, docx_path):

    pdf = fitz.open(pdf_path)
    document = Document()

    first_page = pdf[0]
    page_width_pt = first_page.rect.width
    page_width_in = page_width_pt / POINTS_PER_INCH
    page_height_in = first_page.rect.height / POINTS_PER_INCH

    left_margin_in, right_margin_in, content_bounds = detect_page_margins(pdf, page_width_pt)

    section = document.sections[0]
    section.page_width = Inches(page_width_in)
    section.page_height = Inches(page_height_in)
    section.left_margin = Inches(left_margin_in)
    section.right_margin = Inches(right_margin_in)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    usable_width_in = page_width_in - left_margin_in - right_margin_in

    normal_style = document.styles['Normal']
    normal_style.paragraph_format.space_after = Pt(2)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.0

    enable_hyphenation(document)

    page_bg_color = get_page_background_color(first_page)
    if page_bg_color:
        print(f"[docx_builder] page background detected: {page_bg_color}")
        set_document_background(document, page_bg_color)

    ocr_pages_used = []

    for page_num, page in enumerate(pdf):
        text_dict, page_used_ocr = get_page_text_dict(page)
        if page_used_ocr:
            ocr_pages_used.append(page_num + 1)
            print(f"OCR used for page {page_num + 1} (no native text layer found)")

        if page_used_ocr:
            found_tables = None
            table_bboxes = []
        else:
            found_tables = page.find_tables()
            table_bboxes = [t.bbox for t in found_tables.tables]

        items = []

        if found_tables:
            for table in found_tables.tables:
                items.append({'bbox': table.bbox, 'kind': 'table', 'data': table.extract(), 'table_obj': table})
        print(f"[docx_builder] page {page_num+1}: tables found={len(table_bboxes)}")
        for tb in table_bboxes:
            print(f"  table bbox={tb}")
        horizontal_lines = get_horizontal_lines(page)
        for block in text_dict["blocks"]:
            if block["type"] not in (0, 1):
                continue
            block_bbox = block["bbox"]

            if block["type"] == 0:
                is_inside_table = any(
                    rects_overlap_ratio(block_bbox, tb) >= TABLE_OVERLAP_THRESHOLD
                    for tb in table_bboxes
                )
                if is_inside_table:
                    continue
            elif block["type"] == 1 and page_used_ocr:
                continue

            items.append({'bbox': block_bbox, 'kind': 'block', 'data': block})

        pre_merge_narrow_items = [
            i for i in items
            if (i['bbox'][2] - i['bbox'][0]) < FULL_WIDTH_RATIO_THRESHOLD * page_width_pt
        ]
        column_gap = find_column_gap(pre_merge_narrow_items, page_width_pt)
        print(f"[docx_builder] page {page_num+1}: pre-merge column_gap={column_gap}")

        items = find_row_pairs(items, page_width_pt, column_gap)
        items = merge_same_line_blocks(items, column_gap)

        print(f"[docx_builder] page {page_num+1}: total items={len(items)}")
        for i in items:
            block_bbox = i['bbox']
            preview = ""
            if i['kind'] == 'block' and i['data']['type'] == 0:
                preview = " ".join(
                    span["text"] for line in i['data']["lines"] for span in line["spans"]
                )[:50]
            elif i['kind'] == 'row_pair':
                left_preview = " ".join(
                    span["text"] for line in i['left_data']["lines"] for span in line["spans"]
                )[:30]
                right_preview = " ".join(
                    span["text"] for line in i['right_data']["lines"] for span in line["spans"]
                )[:30]
                preview = f"[ROW_PAIR] {left_preview} || {right_preview}"
            print(f"  item kind={i['kind']} bbox={block_bbox} width={block_bbox[2]-block_bbox[0]:.1f} text='{preview}'")

        narrow_items = [
            i for i in items
            if (i['bbox'][2] - i['bbox'][0]) < FULL_WIDTH_RATIO_THRESHOLD * page_width_pt
        ]
        gap = find_column_gap(narrow_items, page_width_pt)

        if gap:
            gap_start, gap_end = gap
            groups = order_items_for_columns(items, gap_start, gap_end, page_width_pt)

            sidebar_color = get_sidebar_fill_color(page, gap_start)
            print(f"[docx_builder] page {page_num + 1}: sidebar_color detected: {sidebar_color}")

            for item in groups['above']:
                if item['kind'] == 'table':
                    add_table_to_doc(document, item['data'], usable_width_in, page, item.get('table_obj'))
                elif item['kind'] == 'row_pair':
                    add_row_pair_to_container(document, item, content_bounds)
                else:
                    add_block_to_container(document, item['data'], page_width_pt, content_bounds, horizontal_lines)

            render_two_column_table(document, groups['left'], groups['right'],
                                     gap_start, gap_end, page_width_pt,
                                     usable_width_in, sidebar_color)

            for item in groups['within'] + groups['below']:
                if item['kind'] == 'table':
                    add_table_to_doc(document, item['data'], usable_width_in, page, item.get('table_obj'))
                elif item['kind'] == 'row_pair':
                    add_row_pair_to_container(document, item, content_bounds)
                else:
                    add_block_to_container(document, item['data'], page_width_pt, content_bounds, horizontal_lines)

            continue  # this page is done — skip the single-column loop below

        # No column gap found -> single-column fallback
        print(f"[docx_builder] page {page_num + 1}: no column gap found, falling back to single-column y-sort")
        items.sort(key=lambda item: item['bbox'][1])

        for item in items:
            if item['kind'] == 'table':
                add_table_to_doc(document, item['data'], usable_width_in, page, item.get('table_obj'))
            elif item['kind'] == 'row_pair':
                add_row_pair_to_container(document, item, content_bounds)
            else:
                add_block_to_container(document, item['data'], page_width_pt, content_bounds, horizontal_lines)

    document.save(docx_path)
    pdf.close()
    return ocr_pages_used


def add_picture_to_container(container, image_stream, width_in):
    """Document has add_picture() directly, but a table _Cell doesn't —
    for cells, the image must go through a run instead."""
    if hasattr(container, "add_picture"):
        container.add_picture(image_stream, width=Inches(width_in))
    else:
        paragraph = container.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(width_in))


def set_paragraph_bottom_border(paragraph, sz=6, color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def find_line_below(block_bbox, horizontal_lines, min_width_pt):
    bottom_y = block_bbox[3]
    for y, x0, x1, width in horizontal_lines:
        if 0 <= (y - bottom_y) <= 8 and (x1 - x0) >= min_width_pt:
            return True
    return False