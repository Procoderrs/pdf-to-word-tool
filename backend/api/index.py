from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import fitz  # PyMuPDF
import tempfile
import os
import uuid
import io

app = Flask(__name__)
CORS(app)

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

FLAG_ITALIC = 2
FLAG_BOLD = 16
POINTS_PER_INCH = 72

HEADING_SIZE_THRESHOLD = 20  # pt
HEADING_FONT_SIZE = Pt(20)

TABLE_OVERLAP_THRESHOLD = 0.5

# Fallback font if a PDF's embedded font name can't be cleanly mapped —
# every system has this, so text never silently falls back to something
# unpredictable.
DEFAULT_FONT_NAME = "Calibri"


@app.route('/', methods=['GET'])
def index():
    return jsonify({'message': 'API is working'})


def inspect_pdf(path):
    doc = fitz.open(path)
    try:
        encrypted = doc.is_encrypted
        has_text = False
        if not encrypted:
            for page in doc:
                if page.get_text().strip():
                    has_text = True
                    break
        return {'encrypted': encrypted, 'has_text': has_text}
    finally:
        doc.close()


def clean_font_name(raw_font_name):
    """
    PDF embedded font names often look like 'AAAAAA+CanvaSans-Bold'
    (a subset prefix + PostScript style suffix). Word doesn't recognize
    that exact name, so we strip the subset prefix and common style
    suffixes to get something closer to a real font family name. If the
    result still isn't a font installed on the reader's machine, Word
    silently substitutes a default — this just gives it a cleaner name
    to try first.
    """
    if not raw_font_name:
        return DEFAULT_FONT_NAME

    name = raw_font_name
    if '+' in name:
        name = name.split('+', 1)[1]  # drop subset prefix like "AAAAAA+"

    for suffix in ('-Bold', '-Italic', '-BoldItalic', '-Regular', ',Bold', ',Italic'):
        if name.endswith(suffix):
            name = name[: -len(suffix)]

    return name or DEFAULT_FONT_NAME


def rects_overlap_ratio(block_bbox, table_bbox):
    bx0, by0, bx1, by1 = block_bbox
    tx0, ty0, tx1, ty1 = table_bbox

    ix0, iy0 = max(bx0, tx0), max(by0, ty0)
    ix1, iy1 = min(bx1, tx1), min(by1, ty1)

    if ix1 <= ix0 or iy1 <= iy0:
        return 0.0

    intersection_area = (ix1 - ix0) * (iy1 - iy0)
    block_area = (bx1 - bx0) * (by1 - by0)
    if block_area == 0:
        return 0.0

    return intersection_area / block_area


def set_repeat_header_row(row):
    """
    python-docx has no high-level API for 'repeat as header row on every
    page' — it has to be set via the raw OOXML <w:trPr><w:tblHeader/>
    element on the row.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def add_table_to_doc(document, table_data, usable_width_in):
    if not table_data or not table_data[0]:
        return

    num_rows = len(table_data)
    num_cols = len(table_data[0])

    docx_table = document.add_table(rows=num_rows, cols=num_cols)
    docx_table.style = 'Table Grid'

    # Fixed, equal column widths based on the page's usable width — avoids
    # Word auto-compressing columns for wide text (e.g. long English words
    # in a large data table), which is what causes cramped/broken cells.
    docx_table.autofit = False
    col_width = Inches(usable_width_in / num_cols)
    for col in docx_table.columns:
        col.width = col_width

    for row_index, row in enumerate(table_data):
        for col_index, cell_text in enumerate(row):
            if col_index >= num_cols:
                continue
            cell = docx_table.cell(row_index, col_index)
            cell.text = str(cell_text) if cell_text is not None else ''
            cell.width = col_width  # python-docx needs this set per-cell too

        # Repeat the first row as a header on every page this table spans —
        # important for large tables that split across multiple pages.
        if row_index == 0:
            set_repeat_header_row(docx_table.rows[0])


def convert_pdf_to_docx(pdf_path, docx_path):
    pdf = fitz.open(pdf_path)
    document = Document()

    # --- Page size + margins: match the PDF's own page size instead of
    # Word's default Letter/1-inch-margin layout, so text wraps roughly
    # where it did in the source PDF. ---
    first_page = pdf[0]
    page_width_in = first_page.rect.width / POINTS_PER_INCH
    page_height_in = first_page.rect.height / POINTS_PER_INCH

    section = document.sections[0]
    section.page_width = Inches(page_width_in)
    section.page_height = Inches(page_height_in)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    usable_width_in = page_width_in - 2  # minus left+right margin, for table sizing

    # Tighten default paragraph spacing — Word's default "Normal" style adds
    # extra space after each paragraph, which compounds since every PDF
    # line/paragraph becomes its own docx paragraph, making the document
    # visually "looser" than the source PDF.
    normal_style = document.styles['Normal']
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.line_spacing = 1.0

    for page_num, page in enumerate(pdf):
        text_dict = page.get_text("dict")

        found_tables = page.find_tables()
        table_bboxes = [t.bbox for t in found_tables.tables]

        items = []

        for table in found_tables.tables:
            items.append({
                'y': table.bbox[1],
                'kind': 'table',
                'data': table.extract(),
            })

        for block in text_dict["blocks"]:
            if block["type"] not in (0, 1):
                continue

            block_bbox = block["bbox"]

            if block["type"] == 0:
                is_inside_table = any(
                    rects_overlap_ratio(block_bbox, tb) >= TABLE_OVERLAP_THRESHOLD
                    for tb in table_bboxes
                )
                if is_inside_table:
                    continue

            items.append({'y': block_bbox[1], 'kind': 'block', 'data': block})

        items.sort(key=lambda item: item['y'])

        for item in items:
            if item['kind'] == 'table':
                add_table_to_doc(document, item['data'], usable_width_in)
                continue

            block = item['data']

            if block["type"] == 0:
                paragraph = document.add_paragraph()
                max_size_in_paragraph = 0

                for line_index, line in enumerate(block["lines"]):
                    for span in line["spans"]:
                        text = span["text"]
                        if not text.strip():
                            continue

                        run = paragraph.add_run(text)
                        span_size = span["size"]
                        max_size_in_paragraph = max(max_size_in_paragraph, span_size)
                        run.font.size = Pt(round(span_size))
                        run.font.name = clean_font_name(span.get("font", ""))

                        flags = span["flags"]
                        run.bold = bool(flags & FLAG_BOLD)
                        run.italic = bool(flags & FLAG_ITALIC)

                        color_int = span.get("color", 0)
                        r = (color_int >> 16) & 255
                        g = (color_int >> 8) & 255
                        b = color_int & 255
                        run.font.color.rgb = RGBColor(r, g, b)

                    if line_index < len(block["lines"]) - 1:
                        paragraph.add_run(" ")

                if max_size_in_paragraph >= HEADING_SIZE_THRESHOLD:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = HEADING_FONT_SIZE

            elif block["type"] == 1:
                image_bytes = block.get("image")
                if not image_bytes:
                    continue

                bbox = block["bbox"]
                width_pt = bbox[2] - bbox[0]
                width_in = width_pt / POINTS_PER_INCH

                try:
                    image_stream = io.BytesIO(image_bytes)
                    document.add_picture(image_stream, width=Inches(width_in))
                except Exception as e:
                    print(f"Image skip (page {page_num + 1}): {e}")

    document.save(docx_path)
    pdf.close()


@app.route('/api/convert', methods=['POST'])
def convert_pdf_to_word():
    if 'pdfFile' not in request.files:
        return jsonify({'error': 'Please select a valid PDF file.'}), 400

    file = request.files['pdfFile']

    if file.filename == '':
        return jsonify({'error': 'Please select a valid PDF file.'}), 400

    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Only PDF files are allowed.'}), 400

    file_bytes = file.read()

    if len(file_bytes) == 0:
        return jsonify({'error': 'The uploaded file is empty or corrupted.'}), 400

    if len(file_bytes) > MAX_FILE_SIZE:
        return jsonify({'error': 'File exceeds the 50MB limit.'}), 400

    unique_id = uuid.uuid4().hex
    input_path = os.path.join(tempfile.gettempdir(), f'input_{unique_id}.pdf')
    output_path = os.path.join(tempfile.gettempdir(), f'output_{unique_id}.docx')

    try:
        with open(input_path, 'wb') as f:
            f.write(file_bytes)

        try:
            info = inspect_pdf(input_path)
        except Exception as e:
            print('PDF inspection error:', str(e))
            return jsonify({'error': 'This file could not be read as a valid PDF.'}), 400

        if info['encrypted']:
            return jsonify({
                'error': 'This PDF is password-protected. Please remove the password and try again.'
            }), 422

        image_only_warning = not info['has_text']

        convert_pdf_to_docx(input_path, output_path)

        if not os.path.exists(output_path):
            return jsonify({'error': 'Conversion failed. Please try a different file.'}), 500

        with open(output_path, 'rb') as f:
            docx_bytes = f.read()

        original_name = os.path.splitext(file.filename)[0]

        response = send_file(
            io.BytesIO(docx_bytes),
            as_attachment=True,
            download_name=f'{original_name}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )

        if image_only_warning:
            response.headers['X-Conversion-Mode'] = 'image-only'

        return response

    except Exception as e:
        print('Conversion error:', str(e))
        return jsonify({'error': 'Conversion failed. Please try a different file.'}), 500

    finally:
        for path in (input_path, output_path):
            if os.path.exists(path):
                try:
                    os.remove(path)
                except Exception:
                    pass


@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})


if __name__ == '__main__':
    app.run(debug=True, port=5001)